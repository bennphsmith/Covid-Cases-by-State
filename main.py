###### Imports ######

import os
import docx
import yaml
import shutil
import requests

import pandas as pd

from datetime import datetime
from custom_errors import StatusCodeInvalid, ExpectedColumnsNotFound, IncorrectNumberOfDates


###### Collect and Process Data from API ######

def generate_soql_query(
        url: str,
        number_of_days: int,
        state_name: str
) -> str:
    """
    Calculates the number call limit based on number of days requested and generates SoQL query
    Args:
        url: String - based URL for dataset
        number_of_days: Int - Number of days to include in the report
    Returns:
        str - SoQL query used to call the JSON response
    """
    limit = number_of_days*3
    query = f"""?$query=
        SELECT%20%2A%20
        WHERE%20state_name%3D%27{state_name}%27%20
        ORDER%20BY%20date%20DESC%20
        LIMIT%20{limit}"""
    return url+query


def get_testing_data(
        url: str,
) -> str:
    """
    Collects the Data from the URL, checks Status Code validity and loads into a DataFrame
    Args:
        url: String - API URL for get request to collect data
    Returns:
        List - JSON String containing API payload
    """
    response = requests.get(url)
    if response.status_code != 200:
        raise StatusCodeInvalid(url=url, status_code=response.status_code)
    else:
        return response.json()


def convert_testing_data_to_dataframe(
        covid_testing_json: str,
        expected_schema: dict,
) -> pd.DataFrame:
    """
    Converts the JSON payload from the API request in Pandas DataFrame after performing checks
    Args:
        covid_testing_json: List - JSON payload from API call
        expected_schema: List - Expected Columns in the Output
    Returns:
        covid_testing_dataframe: Pandas DataFrame - Contains all info from payload
    """
    df_covid_testing = pd.DataFrame(covid_testing_json)
    missing_columns = [x for x in list(expected_schema.keys()) if x not in list(df_covid_testing.columns)]
    if missing_columns:
        raise ExpectedColumnsNotFound(missing_columns)
    else:
        return df_covid_testing.astype(expected_schema)


###### Merge DataFrame Outputs ######

def create_output_table(
        df_covid_testing: pd.DataFrame,
        outcome_type: str,
        number_of_days: int,
        output_schema: dict
) -> pd.DataFrame:
    """
    Covert columns names to those specified in the config Schema and returns columns selected DataFrame
    Args:
        df_covid_testing: DataFrame - covid testing payload derived from input API
        outcome_type: String - 'Positive' or 'Negative' depending on test outcome
        number_of_days: int - Number of days included in output table
        output_schema: dict - Converts api column names to output column names
    Returns:
        DataFrame - Table in Output Schema
    """
    df_final_output = df_covid_testing.loc[df_covid_testing['overall_outcome'] == outcome_type]
    if len(df_final_output) != number_of_days:
        raise IncorrectNumberOfDates(number_of_days)
    else:
        return df_final_output.rename(columns=output_schema)[list(output_schema.values())]


###### Create Word Doc ######

def create_word_document(
        state_name: str,
        df: pd.DataFrame,
        current_date: str
) -> docx.Document:
    """
    Creates a Word Document with variable Heading and table which replicates the input DataFrame
    Args:
        state: String - state that the DataFrame relates to
        df: DataFrame - Output DataFrame
        current_date: datetime - Current Date
    Returns:
        Docx File - containing data from output DataFrame
    """
    doc = docx.Document()
    doc.add_heading(f"{state_name} - {current_date}")
    t = doc.add_table(df.shape[0] + 1, df.shape[1])
    # Header rows for Table
    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]
    # Data for Table
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = str(df.values[i, j])
    return doc


###### Main ######

if __name__ == "__main__":
    # Generate variables
    config = yaml.safe_load(open("Config/config.yaml"))
    current_date = datetime.today().strftime('%d-%m-%Y')
    zip_file_name = f"{config['zip_file_name']}_{config['number_of_days']}days_{current_date}"

    # Check if Results path is created - if not then create it
    if not os.path.isdir(config['results_directory']):
        os.mkdir(config['results_directory'])

    # Iterate through the State_list and create an Docx file for each API payload
    for state_name in config['state_names']:
        api_query = generate_soql_query(
            url=config['url'],
            number_of_days=config['number_of_days'],
            state_name=state_name
        )
        covid_testing_json = get_testing_data(url=api_query)
        df_covid_testing = convert_testing_data_to_dataframe(
            covid_testing_json=covid_testing_json,
            expected_schema=config['expected_schema']
        )
        df_positive = create_output_table(
            df_covid_testing=df_covid_testing,
            outcome_type="Positive",
            number_of_days=config['number_of_days'],
            output_schema=config['positive_table_schema']
        )
        df_negative = create_output_table(
            df_covid_testing=df_covid_testing,
            outcome_type="Negative",
            number_of_days=config['number_of_days'],
            output_schema=config['negative_table_schema']
        )
        df_final_output = pd.merge(
            df_positive,
            df_negative,
            how='left',
            on='Date'
        )[config['output_columns']]
        word_doc = create_word_document(
            state_name=state_name,
            df=df_final_output,
            current_date=current_date
        )

        # Save Doc file to Results folder
        doc_name = f"{config['results_directory']}/{state_name}-{current_date}.docx"
        word_doc.save(doc_name)
        print(f"Successfully saved: {doc_name}")

    # Make Zip folder from Results folder
    shutil.make_archive(zip_file_name, 'zip', config['results_directory'])
